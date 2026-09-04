import codecs
import re
import zipfile
from dataclasses import dataclass
from io import BytesIO

from shared.config import get_settings
from shared.files.errors import FileTranslationError

SUPPORTED_EXTENSIONS = (".txt", ".pdf", ".docx", ".srt")

MAX_UNITS = 20_000
MAX_PDF_PAGES = 500
MAX_UNCOMPRESSED_BYTES = 80 * 1024 * 1024
MAX_COMPRESSION_RATIO = 150
MAX_ZIP_ENTRIES = 2000

TOO_MUCH_TEXT_MESSAGE = (
	"There's too much text in that file for me to translate in one go. "
	"Split it into smaller files and send them one at a time."
)
SUSPICIOUS_ARCHIVE_MESSAGE = (
	"That file unpacks to far more than its size suggests, so I'm not opening it."
)

PARAGRAPH_SPLIT_RE = re.compile(r"\n\s*\n")

BOM_ENCODINGS = (
	(codecs.BOM_UTF8, "utf-8-sig"),
	(codecs.BOM_UTF32_LE, "utf-32"),
	(codecs.BOM_UTF32_BE, "utf-32"),
	(codecs.BOM_UTF16_LE, "utf-16"),
	(codecs.BOM_UTF16_BE, "utf-16"),
)
FALLBACK_ENCODINGS = ("utf-8", "cp1252")

UNREADABLE_TEXT_MESSAGE = (
	"I couldn't work out the character encoding of that file. Save it as UTF-8 and send it again."
)


def max_extracted_chars() -> int:
	return get_settings().max_extracted_chars


def decode_text(data: bytes) -> str:
	"""Decode a plain-text file without silently turning unknown bytes into mojibake."""
	for bom, encoding in BOM_ENCODINGS:
		if data.startswith(bom):
			try:
				return data.decode(encoding)
			except UnicodeDecodeError as exc:
				raise FileTranslationError(UNREADABLE_TEXT_MESSAGE) from exc

	for encoding in FALLBACK_ENCODINGS:
		try:
			return data.decode(encoding)
		except UnicodeDecodeError:
			continue
	raise FileTranslationError(UNREADABLE_TEXT_MESSAGE)


@dataclass
class ExtractedContent:
	units: list[str]
	rebuild_context: object | None = None


def extract_content(data: bytes, extension: str) -> ExtractedContent:
	if extension == ".txt":
		return extract_txt(data)
	if extension == ".pdf":
		return extract_pdf(data)
	if extension == ".docx":
		return extract_docx(data)
	if extension == ".srt":
		return extract_srt(data)
	raise FileTranslationError(f"Unsupported extension: {extension}")


def split_paragraphs(text: str) -> list[str]:
	return [block.strip() for block in PARAGRAPH_SPLIT_RE.split(text) if block.strip()]


def guard_text_size(text: str) -> str:
	if len(text) > max_extracted_chars():
		raise FileTranslationError(TOO_MUCH_TEXT_MESSAGE)
	return text


def guard_units(units: list[str]) -> list[str]:
	if len(units) > MAX_UNITS:
		raise FileTranslationError(TOO_MUCH_TEXT_MESSAGE)
	guard_text_size("".join(units))
	return units


def guard_archive(data: bytes) -> None:
	"""Refuse archives that inflate far beyond their packed size, before any parser sees them."""
	try:
		with zipfile.ZipFile(BytesIO(data)) as archive:
			entries = archive.infolist()
			if len(entries) > MAX_ZIP_ENTRIES:
				raise FileTranslationError(SUSPICIOUS_ARCHIVE_MESSAGE)
			uncompressed = sum(entry.file_size for entry in entries)
	except FileTranslationError:
		raise
	except zipfile.BadZipFile as exc:
		raise FileTranslationError("Could not read that .docx file, it might be corrupted.") from exc

	if uncompressed > MAX_UNCOMPRESSED_BYTES:
		raise FileTranslationError(SUSPICIOUS_ARCHIVE_MESSAGE)
	if data and uncompressed / len(data) > MAX_COMPRESSION_RATIO:
		raise FileTranslationError(SUSPICIOUS_ARCHIVE_MESSAGE)


def extract_txt(data: bytes) -> ExtractedContent:
	if len(data) > max_extracted_chars() * 4:
		raise FileTranslationError(TOO_MUCH_TEXT_MESSAGE)
	text = guard_text_size(decode_text(data).strip())
	if not text:
		raise FileTranslationError("That file has no text to translate.")
	return ExtractedContent(units=guard_units(split_paragraphs(text) or [text]))


def extract_pdf(data: bytes) -> ExtractedContent:
	from pypdf import PdfReader

	try:
		reader = PdfReader(BytesIO(data))
		if reader.is_encrypted:
			try:
				unlocked = reader.decrypt("")
			except Exception:
				unlocked = 0
			if not unlocked:
				raise FileTranslationError("That PDF is password-protected, I can't read it.")

		if len(reader.pages) > MAX_PDF_PAGES:
			raise FileTranslationError(
				f"That PDF has too many pages, I can handle up to {MAX_PDF_PAGES} at a time."
			)

		pages_text: list[str] = []
		total = 0
		for page in reader.pages:
			page_text = page.extract_text() or ""
			total += len(page_text)
			if total > max_extracted_chars():
				raise FileTranslationError(TOO_MUCH_TEXT_MESSAGE)
			pages_text.append(page_text)
	except FileTranslationError:
		raise
	except Exception as exc:
		raise FileTranslationError("Could not read that PDF, it might be corrupted or protected.") from exc

	text = "\n\n".join(page.strip() for page in pages_text if page.strip())
	if not text:
		raise FileTranslationError("Couldn't find any text in that PDF.")
	return ExtractedContent(units=guard_units(split_paragraphs(text) or [text]))


def table_paragraphs(table) -> list[str]:
	"""Every paragraph inside a table, including nested ones. A merged cell is read once."""
	found = []
	seen = set()
	for row in table.rows:
		for cell in row.cells:
			if id(cell._tc) in seen:
				continue
			seen.add(id(cell._tc))
			found.extend(paragraph.text.strip() for paragraph in cell.paragraphs)
			for nested in cell.tables:
				found.extend(table_paragraphs(nested))
	return found


def extract_docx(data: bytes) -> ExtractedContent:
	from docx import Document

	guard_archive(data)

	try:
		document = Document(BytesIO(data))
		paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
		for table in document.tables:
			paragraphs.extend(table_paragraphs(table))
		for section in document.sections:
			for part in (section.header, section.footer):
				paragraphs.extend(paragraph.text.strip() for paragraph in part.paragraphs)
	except FileTranslationError:
		raise
	except Exception as exc:
		raise FileTranslationError("Could not read that .docx file, it might be corrupted.") from exc

	units = [paragraph for paragraph in paragraphs if paragraph]
	if not units:
		raise FileTranslationError("Couldn't find any text in that document.")
	return ExtractedContent(units=guard_units(units))


def extract_srt(data: bytes) -> ExtractedContent:
	import srt as srt_lib

	if len(data) > max_extracted_chars() * 4:
		raise FileTranslationError(TOO_MUCH_TEXT_MESSAGE)

	try:
		raw = guard_text_size(decode_text(data))
		subtitles = list(srt_lib.parse(raw))
	except FileTranslationError:
		raise
	except Exception as exc:
		raise FileTranslationError("Could not read that subtitle file, it might be corrupted.") from exc

	if not subtitles:
		raise FileTranslationError("Couldn't find any subtitles in that file.")

	units = [subtitle.content for subtitle in subtitles]
	return ExtractedContent(units=guard_units(units), rebuild_context=subtitles)
